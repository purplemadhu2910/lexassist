from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from contextlib import asynccontextmanager
from pydantic import BaseModel, field_validator
from ai_engine import AIEngine
from document_parser import DocumentParser
from database import Database
from rag_engine import preload_resources
import csv
import io
import logging
import os
import time
from collections import defaultdict
from dotenv import load_dotenv

load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), "..", ".env"))

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("Starting up — pre-loading RAG resources...")
    preload_resources()
    yield


app = FastAPI(
    title="LexAssist API",
    description="AI-powered legal and tax assistant backend",
    version="1.0.0",
    lifespan=lifespan
)

_frontend_url = os.getenv("FRONTEND_URL", "")
ALLOWED_ORIGINS = ["http://localhost:8501", "http://127.0.0.1:8501"]
if _frontend_url:
    ALLOWED_ORIGINS += [u.strip() for u in _frontend_url.split(",")]
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["GET", "POST", "PATCH", "DELETE"],
    allow_headers=["*"],
)

ai_engine = AIEngine()
doc_parser = DocumentParser()
db = Database()

# ── Rate limiter ──────────────────────────────────────────────────────────
_login_attempts: dict = defaultdict(list)
_ask_attempts: dict = defaultdict(list)
MAX_LOGIN_ATTEMPTS = 5
MAX_ASK_ATTEMPTS = 20
RATE_LIMIT_WINDOW = 60  # seconds

def check_rate_limit(ip: str):
    now = time.time()
    _login_attempts[ip] = [t for t in _login_attempts[ip] if now - t < RATE_LIMIT_WINDOW]
    if len(_login_attempts[ip]) >= MAX_LOGIN_ATTEMPTS:
        raise HTTPException(status_code=429, detail="Too many login attempts. Try again in a minute.")
    _login_attempts[ip].append(now)

def check_ask_rate_limit(user_id: int):
    now = time.time()
    _ask_attempts[user_id] = [t for t in _ask_attempts[user_id] if now - t < RATE_LIMIT_WINDOW]
    if len(_ask_attempts[user_id]) >= MAX_ASK_ATTEMPTS:
        raise HTTPException(status_code=429, detail="Too many requests. Please wait a minute before asking again.")
    _ask_attempts[user_id].append(now)

# ── Auth ──────────────────────────────────────────────────────────────────

def get_current_user(request: Request) -> int:
    token = request.headers.get("X-Auth-Token")
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id = db.get_session_user(token)
    if user_id is None:
        raise HTTPException(status_code=401, detail="Session expired or invalid. Please log in again.")
    db.refresh_session(token)
    return user_id

def get_token(request: Request) -> str:
    token = request.headers.get("X-Auth-Token")
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return token

# ── Input sanitization ────────────────────────────────────────────────────
BLOCKED_PATTERNS = [
    "ignore previous", "ignore prior", "ignore all", "disregard",
    "you are now", "you are a", "act as", "pretend to be", "pretend you are",
    "new instructions", "system prompt", "forget everything", "forget your",
    "override", "jailbreak", "do anything now", "dan mode",
    "bypass", "ignore your instructions", "ignore the above",
    "from now on", "your new role", "your true self",
]

def sanitize_query(query: str) -> str:
    lower = query.lower()
    for pattern in BLOCKED_PATTERNS:
        if pattern in lower:
            raise HTTPException(status_code=400, detail="Invalid query content.")
    return query.strip()[:2000]

# ── Models ────────────────────────────────────────────────────────────────

class QueryRequest(BaseModel):
    query: str
    category: str = "general"
    history: list = []
    language: str = "English"

    @field_validator("category")
    @classmethod
    def validate_category(cls, v):
        if v not in ("legal", "tax", "general", "document"):
            return "general"
        return v

    @field_validator("language")
    @classmethod
    def validate_language(cls, v):
        allowed = {"English", "Hindi", "Tamil", "Telugu", "Kannada", "Malayalam", "Bengali", "Marathi", "Gujarati"}
        return v if v in allowed else "English"

class QueryResponse(BaseModel):
    response: str
    category: str
    suggested_questions: list[str] = []
    sources: list[str] = []

class AuthRequest(BaseModel):
    username: str
    password: str

    @field_validator("username")
    @classmethod
    def validate_username(cls, v):
        v = v.strip()
        if not v or len(v) > 50:
            raise ValueError("Invalid username")
        return v

    @field_validator("password")
    @classmethod
    def validate_password(cls, v):
        if not v or len(v) < 6 or len(v) > 128:
            raise ValueError("Password must be 6–128 characters")
        if not any(c.isdigit() or not c.isalpha() for c in v):
            raise ValueError("Password must contain at least one number or special character")
        return v

class BookmarkRequest(BaseModel):
    query_id: int

class CaseLawRequest(BaseModel):
    query: str

class DraftDocumentRequest(BaseModel):
    doc_type: str
    details: dict = {}

class SectionLookupRequest(BaseModel):
    act: str
    section: str

class TimelineRequest(BaseModel):
    situation: str

class PenaltyRequest(BaseModel):
    section: str
    circumstances: str = ""

class GlossaryRequest(BaseModel):
    term: str

class BookmarkNoteRequest(BaseModel):
    note: str = ""

    @field_validator("note")
    @classmethod
    def trim_note(cls, v):
        return v.strip()[:500]

class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str

    @field_validator("new_password")
    @classmethod
    def validate_new_password(cls, v):
        if not v or len(v) < 6 or len(v) > 128:
            raise ValueError("New password must be 6–128 characters")
        return v

class FeedbackRequest(BaseModel):
    query_id: int
    rating: int  # 1 (thumbs up) or -1 (thumbs down)
    comment: str = ""

class ExportRequest(BaseModel):
    title: str = "LexAssist Document"
    content: str

# ── Routes ────────────────────────────────────────────────────────────────

@app.get("/")
async def root():
    return {"message": "Welcome to LexAssist API", "version": "1.0.0"}

@app.get("/health")
async def health_check():
    try:
        stats = db.get_stats()
        return {
            "status": "healthy",
            "database": "connected",
            "total_queries": stats["total_queries"],
            "queries_by_category": stats["by_category"]
        }
    except Exception as e:
        return {"status": "degraded", "error": str(e)}

@app.post("/register")
async def register(request: AuthRequest):
    success = db.register_user(request.username, request.password)
    if not success:
        raise HTTPException(status_code=400, detail="Username already exists")
    return {"message": "Account created successfully"}

@app.post("/login")
async def login(request: AuthRequest, req: Request):
    client_ip = req.client.host
    check_rate_limit(client_ip)

    user = db.login_user(request.username, request.password)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid username or password")

    # Purge stale sessions opportunistically
    db.purge_expired_sessions()

    token = db.create_session(user["id"])
    return {"message": "Login successful", "token": token, "user_id": user["id"], "username": user["username"]}

@app.post("/logout")
async def logout(token: str = Depends(get_token)):
    db.delete_session(token)
    return {"message": "Logged out successfully"}

@app.post("/ask", response_model=QueryResponse)
async def ask_question(request: QueryRequest, req: Request, user_id: int = Depends(get_current_user)):
    check_ask_rate_limit(user_id)
    try:
        query = sanitize_query(request.query)
        logger.info(f"Processing query: {query[:50]}...")
        response, sources = await ai_engine.process_query(query, request.category, request.history, request.language)
        db.save_query(query, response, request.category, user_id)
        suggestions = ai_engine.generate_suggestions(query, request.category)
        return QueryResponse(
            response=response,
            category=request.category,
            suggested_questions=suggestions,
            sources=sources
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing query: {str(e)}")
        raise HTTPException(status_code=500, detail="Error processing query")

@app.post("/explain-document")
async def explain_document(file: UploadFile = File(...), user_id: int = Depends(get_current_user)):
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")

        file_ext = '.' + file.filename.split('.')[-1].lower()
        allowed_extensions = ['.pdf', '.txt', '.docx']
        if file_ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}")

        content = await file.read()
        if len(content) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 5MB.")

        extracted_text = doc_parser.extract_text(content, file_ext)
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from document")

        word_count = len(extracted_text.split())
        reading_time = max(1, round(word_count / 200))  # avg 200 wpm
        explanation = await ai_engine.explain_document(extracted_text)
        return {
            "filename": file.filename,
            "extracted_text": extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text,
            "explanation": explanation,
            "text_length": len(extracted_text),
            "word_count": word_count,
            "reading_time_minutes": reading_time,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise HTTPException(status_code=500, detail="Error processing document")

@app.get("/history")
async def get_history(limit: int = 10, offset: int = 0, search: str = "", category: str = "", user_id: int = Depends(get_current_user)):
    limit = min(max(1, limit), 100)
    offset = max(0, offset)
    search_term = search.strip()[:200] if search else None
    category_filter = category.strip() if category and category != "All" else None
    try:
        history = db.get_history(limit, offset, user_id, search_term, category_filter)
        total = db.get_history_count(user_id, search_term, category_filter)
        return {"history": history, "count": len(history), "total": total}
    except Exception as e:
        logger.error(f"Error fetching history: {str(e)}")
        raise HTTPException(status_code=500, detail="Error fetching history")

@app.get("/history/export")
async def export_history(user_id: int = Depends(get_current_user)):
    try:
        history = db.get_history(limit=10000, offset=0, user_id=user_id)
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=["id", "timestamp", "category", "query", "response"])
        writer.writeheader()
        for item in history:
            writer.writerow({
                "id": item["id"],
                "timestamp": item["timestamp"],
                "category": item["category"],
                "query": item["query"],
                "response": item["response"]
            })
        output.seek(0)
        return StreamingResponse(
            iter([output.getvalue()]),
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename=lexassist_history.csv"}
        )
    except Exception as e:
        logger.error(f"Error exporting history: {str(e)}")
        raise HTTPException(status_code=500, detail="Error exporting history")

@app.post("/bookmarks/toggle")
async def toggle_bookmark(request: BookmarkRequest, user_id: int = Depends(get_current_user)):
    try:
        bookmarked = db.toggle_bookmark(user_id, request.query_id)
        return {"bookmarked": bookmarked, "query_id": request.query_id}
    except Exception as e:
        logger.error(f"Error toggling bookmark: {str(e)}")
        raise HTTPException(status_code=500, detail="Error toggling bookmark")

@app.get("/bookmarks")
async def get_bookmarks(user_id: int = Depends(get_current_user)):
    try:
        bookmarks = db.get_bookmarks(user_id)
        return {"bookmarks": bookmarks, "count": len(bookmarks)}
    except Exception as e:
        logger.error(f"Error fetching bookmarks: {str(e)}")
        raise HTTPException(status_code=500, detail="Error fetching bookmarks")

@app.patch("/bookmarks/{query_id}/note")
async def update_bookmark_note(query_id: int, request: BookmarkNoteRequest, user_id: int = Depends(get_current_user)):
    updated = db.update_bookmark_note(user_id, query_id, request.note)
    if not updated:
        raise HTTPException(status_code=404, detail="Bookmark not found")
    return {"message": "Note updated"}

@app.get("/stats")
async def get_user_stats(user_id: int = Depends(get_current_user)):
    try:
        return db.get_user_stats(user_id)
    except Exception as e:
        logger.error(f"Error fetching stats: {str(e)}")
        raise HTTPException(status_code=500, detail="Error fetching stats")

@app.delete("/history/{entry_id}")
async def delete_history_entry(entry_id: int, user_id: int = Depends(get_current_user)):
    deleted = db.delete_history_entry(entry_id, user_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Entry not found")
    return {"message": "Entry deleted"}

@app.post("/change-password")
async def change_password(request: ChangePasswordRequest, user_id: int = Depends(get_current_user)):
    success = db.change_password(user_id, request.old_password, request.new_password)
    if not success:
        raise HTTPException(status_code=400, detail="Current password is incorrect")
    return {"message": "Password changed successfully"}

@app.post("/compare-contracts")
async def compare_contracts(
    file1: UploadFile = File(...),
    file2: UploadFile = File(...),
    user_id: int = Depends(get_current_user)
):
    try:
        texts = []
        for f in (file1, file2):
            if not f.filename:
                raise HTTPException(status_code=400, detail="Both files are required")
            ext = '.' + f.filename.split('.')[-1].lower()
            if ext not in ['.pdf', '.txt', '.docx']:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {f.filename}")
            content = await f.read()
            if len(content) > 5 * 1024 * 1024:
                raise HTTPException(status_code=400, detail=f"{f.filename} exceeds 5MB limit")
            texts.append(doc_parser.extract_text(content, ext))
        result = await ai_engine.compare_contracts(texts[0], texts[1])
        return {"file1": file1.filename, "file2": file2.filename, "comparison": result}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error comparing contracts: {str(e)}")
        raise HTTPException(status_code=500, detail="Error comparing contracts")

@app.post("/case-law-search")
async def case_law_search(request: CaseLawRequest, user_id: int = Depends(get_current_user)):
    try:
        query = sanitize_query(request.query)
        result = await ai_engine.case_law_search(query)
        db.save_query(f"[Case Law] {query}", str(result), "legal", user_id)
        return result
    except Exception as e:
        logger.error(f"Case law search error: {str(e)}")
        raise HTTPException(status_code=500, detail="Error searching case law")

@app.post("/draft-document")
async def draft_document(request: DraftDocumentRequest, user_id: int = Depends(get_current_user)):
    try:
        result = await ai_engine.draft_document(request.doc_type, request.details)
        db.save_query(f"[Draft] {request.doc_type}", result, "document", user_id)
        return {"draft": result, "doc_type": request.doc_type}
    except Exception as e:
        logger.error(f"Document drafting error: {str(e)}")
        raise HTTPException(status_code=500, detail="Error drafting document")

@app.post("/section-lookup")
async def section_lookup(request: SectionLookupRequest, user_id: int = Depends(get_current_user)):
    try:
        result = await ai_engine.lookup_section(request.act, request.section)
        db.save_query(f"[Section] {request.act} S.{request.section}", str(result), "legal", user_id)
        return result
    except Exception as e:
        logger.error(f"Section lookup error: {str(e)}")
        raise HTTPException(status_code=500, detail="Error looking up section")

@app.post("/legal-timeline")
async def legal_timeline(request: TimelineRequest, user_id: int = Depends(get_current_user)):
    try:
        situation = sanitize_query(request.situation)
        result = await ai_engine.build_legal_timeline(situation)
        db.save_query(f"[Timeline] {situation}", str(result), "legal", user_id)
        return result
    except Exception as e:
        logger.error(f"Timeline error: {str(e)}")
        raise HTTPException(status_code=500, detail="Error building timeline")

@app.post("/penalty-calculator")
async def penalty_calculator(request: PenaltyRequest, user_id: int = Depends(get_current_user)):
    try:
        result = await ai_engine.penalty_calculator(request.section, request.circumstances)
        db.save_query(f"[Penalty] IPC/BNS {request.section}", str(result), "legal", user_id)
        return result
    except Exception as e:
        logger.error(f"Penalty calculator error: {str(e)}")
        raise HTTPException(status_code=500, detail="Error calculating penalty")

@app.post("/glossary")
async def glossary_lookup(request: GlossaryRequest, user_id: int = Depends(get_current_user)):
    try:
        result = await ai_engine.glossary_lookup(request.term)
        return result
    except Exception as e:
        logger.error(f"Glossary error: {str(e)}")
        raise HTTPException(status_code=500, detail="Error looking up term")

@app.post("/analyze-contract")
async def analyze_contract(file: UploadFile = File(...), user_id: int = Depends(get_current_user)):
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        file_ext = '.' + file.filename.split('.')[-1].lower()
        if file_ext not in ['.pdf', '.txt', '.docx']:
            raise HTTPException(status_code=400, detail="Unsupported file type. Allowed: pdf, txt, docx")
        content = await file.read()
        if len(content) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 5MB.")
        extracted_text = doc_parser.extract_text(content, file_ext)
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from document")
        result = await ai_engine.analyze_contract_risks(extracted_text)
        db.save_query(f"[Contract Risk Analysis] {file.filename}", str(result), "document", user_id)
        return {"filename": file.filename, "analysis": result}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing contract: {str(e)}")
        raise HTTPException(status_code=500, detail="Error analyzing contract")

@app.post("/ask-stream")
async def ask_question_stream(request: QueryRequest, req: Request, user_id: int = Depends(get_current_user)):
    check_ask_rate_limit(user_id)
    try:
        query = sanitize_query(request.query)
        generator = ai_engine.process_query_stream(query, request.category, request.history, request.language)
        return StreamingResponse(generator, media_type="text/plain")
    except Exception as e:
        logger.error(f"Error streaming query: {str(e)}")
        raise HTTPException(status_code=500, detail="Error streaming query")

@app.post("/feedback")
async def submit_feedback(request: FeedbackRequest, user_id: int = Depends(get_current_user)):
    try:
        db.save_feedback(user_id, request.query_id, request.rating, request.comment)
        return {"message": "Feedback submitted successfully"}
    except Exception as e:
        logger.error(f"Error saving feedback: {str(e)}")
        raise HTTPException(status_code=500, detail="Error saving feedback")

@app.get("/admin/analytics")
async def get_admin_analytics(user_id: int = Depends(get_current_user)):
    try:
        return db.get_admin_analytics()
    except Exception as e:
        logger.error(f"Error fetching admin analytics: {str(e)}")
        raise HTTPException(status_code=500, detail="Error fetching admin analytics")

@app.post("/export/docx")
async def export_docx(request: ExportRequest, user_id: int = Depends(get_current_user)):
    try:
        from docx import Document
        doc = Document()
        
        # Add title
        heading = doc.add_heading(request.title, level=1)
        
        # Add body content
        for line in request.content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
                
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_title = "".join([c if c.isalnum() else "_" for c in request.title[:30]])
        filename = f"{safe_title}.docx"
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        logger.error(f"Error exporting DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail="Error generating DOCX")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
